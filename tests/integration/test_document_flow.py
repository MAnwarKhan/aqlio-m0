from __future__ import annotations

import io
from dataclasses import replace

import pytest
from docx import Document

from app.application.errors import PreparationError, ValidationError
from app.domain import AssetStatus
from tests.helpers import build_service, fixture_bytes


def test_txt_upload_preparation_is_scoped_and_idempotent() -> None:
    service = build_service()
    project = service.create_project("Handbook")
    content = fixture_bytes("employee_handbook.txt")

    first = service.upload_document(project.id, "..\\employee_handbook.txt", content)
    duplicate = service.upload_document(project.id, "another-name.txt", content)
    prepared = service.prepare_document(project.id, first.id)
    prepared_again = service.prepare_document(project.id, first.id)

    assert first.safe_name != first.original_name
    assert first.original_name == "employee_handbook.txt"
    assert duplicate.id == first.id
    assert prepared.status is AssetStatus.READY
    assert prepared_again.id == first.id
    assert len(service.repository.list_assets(project.id)) == 1
    assert service.repository.version_count(project.id) == 1


def test_docx_real_text_extraction() -> None:
    document = Document()
    document.add_paragraph("The travel allowance is available after approval.")
    buffer = io.BytesIO()
    document.save(buffer)
    service = build_service()
    project = service.create_project("Travel")

    asset = service.upload_document(project.id, "travel.docx", buffer.getvalue())
    prepared = service.prepare_document(project.id, asset.id)

    assert prepared.status is AssetStatus.READY
    assert "travel allowance" in (prepared.normalized_text or "")


def test_participant_add_action_uploads_and_prepares_docx() -> None:
    document = Document()
    document.add_paragraph("Annual leave is available after ninety days.")
    buffer = io.BytesIO()
    document.save(buffer)
    service = build_service()
    project = service.create_project("One-step document")

    prepared = service.add_and_prepare_document(
        project.id, "employee_handbook.docx", buffer.getvalue()
    )
    refreshed = service.get_my_project(project.id)

    assert prepared.status is AssetStatus.READY
    assert refreshed.valid_document_count == 1
    assert refreshed.prepared_document_count == 1
    assert refreshed.current_version_id is not None
    assert service.ask_question(project.id, "When is annual leave available?").citations


@pytest.mark.parametrize(
    ("filename", "content", "message"),
    [
        ("script.exe", b"data", "isn't supported"),
        ("fake.pdf", b"not a pdf", "valid document"),
        ("fake.docx", b"not a docx", "valid document"),
    ],
)
def test_unsupported_and_spoofed_files_fail_safely(
    filename: str, content: bytes, message: str
) -> None:
    service = build_service()
    project = service.create_project("Files")

    with pytest.raises(ValidationError, match=message):
        service.upload_document(project.id, filename, content)


def test_empty_text_does_not_become_ready_and_retry_is_safe() -> None:
    service = build_service()
    project = service.create_project("Empty")
    asset = service.upload_document(project.id, "empty.txt", b"   \n   ")

    with pytest.raises(PreparationError, match="usable text"):
        service.prepare_document(project.id, asset.id)
    with pytest.raises(PreparationError, match="usable text"):
        service.prepare_document(project.id, asset.id)

    failed = service.repository.get_asset(asset.id)
    assert failed is not None and failed.status is AssetStatus.FAILED
    assert len(service.repository.list_assets(project.id)) == 1
    assert service.repository.version_count(project.id) == 0


def test_oversized_document_is_rejected_before_storage() -> None:
    service = build_service()
    service.settings = replace(service.settings, max_file_size_mb=1)
    project = service.create_project("Large File")

    with pytest.raises(ValidationError, match="too large"):
        service.upload_document(project.id, "large.txt", b"x" * (1024 * 1024 + 1))

    assert service.repository.list_assets(project.id) == []


def test_storage_failure_during_preparation_records_a_safe_failure() -> None:
    service = build_service()
    project = service.create_project("Unavailable file")
    asset = service.upload_document(project.id, "handbook.txt", b"Useful policy text")
    service.storage.delete(
        workspace_id=project.workspace_id,
        project_id=project.id,
        storage_key=asset.storage_key,
    )

    with pytest.raises(PreparationError, match="couldn't access"):
        service.prepare_document(project.id, asset.id)

    failed = service.repository.get_asset(asset.id)
    assert failed is not None and failed.status is AssetStatus.FAILED
    assert failed.normalized_text is None
